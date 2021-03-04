# -*- coding: utf-8 -*-
"""
Created on Mon Nov 11 22:24:50 2019

@author: Minzel
"""

from docx import Document
# 图片处理
from docx.shared import Inches
# 打开或处理现有文档
document = Document()

# 添加一个段落
paragraph = document.add_paragraph('Lorem ipsum dolor sit amet.')
# 在paragraph之前添加一个prior_paragraph的段落
prior_paragraph = paragraph.insert_paragraph_before('Lorem ipsum')
# 段落之后添加内容
paragraph = document.add_paragraph('Lorem ipsum ')
paragraph.add_run('dolor sit amet.')

# 添加一个标题
document.add_heading('The REAL meaning of the universe')
# 可以选择level 1 to 9 的标题
document.add_heading('The role of dolphins', level=2)

# 添加分页
document.add_page_break()

# 添加表格
table = document.add_table(rows=2, cols=2)
# 访问单元格 0 和 1 表示行 & 列
cell = table.cell(0, 1)
# 单元格添加文本
cell.text = 'parrot, possibly dead'
# 访问一行中的某列单元格
row = table.rows[1]
row.cells[0].text = 'Foo bar to you.'
row.cells[1].text = 'And a hearty foo bar to you too sir!'
# 添加一行
row = table.add_row()

# 单元格迭代
for row in table.rows:
    for cell in row.cells:
        print(cell.text)

# 对行 & 列 进行计数
row_count = len(table.rows)
col_count = len(table.columns)


# get table data -------------
items = (
    (7, '1024', 'Plush kittens'),
    (3, '2042', 'Furbees'),
    (1, '1288', 'French Poodle Collars, Deluxe'),
)

# add table ------------------
table = document.add_table(1, 3)

# populate header row --------
heading_cells = table.rows[0].cells
heading_cells[0].text = 'Qty'
heading_cells[1].text = 'SKU'
heading_cells[2].text = 'Description'

# add a data row for each item
for item in items:
    cells = table.add_row().cells
    cells[0].text = str(item[0])
    cells[1].text = item[1]
    cells[2].text = item[2]

# 选择表格样式
table.style = 'LightShading-Accent1'
# 添加图片并设置宽度
document.add_picture('sonic-the-hedgehog.jpg', width=Inches(1.0))

# 应用段落样式
document.add_paragraph('Lorem ipsum dolor sit amet.', style='ListBullet')
# 效果同上
# 通常，样式名称与在Word用户界面（UI）中显示的样式名称完全相同。
paragraph = document.add_paragraph('Lorem ipsum dolor sit amet.')
paragraph.style = 'List Bullet'

# 字体加粗
paragraph = document.add_paragraph('Lorem ipsum ')
run = paragraph.add_run('dolor')
run.bold = True
paragraph.add_run(' sit amet.')
# 字体样式
paragraph = document.add_paragraph('Normal text, ')
paragraph.add_run('text with emphasis.', 'Emphasis')
# 同上
paragraph = document.add_paragraph('Normal text, ')
run = paragraph.add_run('text with emphasis.')
run.style = 'Emphasis'

# 保存
document.save('first_py_docx_file.docx')